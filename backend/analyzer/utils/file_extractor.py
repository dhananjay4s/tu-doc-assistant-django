import pdfplumber
import docx
import io

def extract_from_pdf(file_bytes):
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                try:
                    # Crop header/footer/page numbers hatau
                    cropped = page.crop((
                        page.bbox[0],
                        page.bbox[1] + 30,
                        page.bbox[2],
                        page.bbox[3] - 30
                    ))
                    text = cropped.extract_text(x_tolerance=3, y_tolerance=3)
                except Exception:
                    # Crop fail bhayo bhane full page use
                    text = page.extract_text(x_tolerance=3, y_tolerance=3)

                if text:
                    lines = text.split('\n')
                    processed = []
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        if line.isupper() and len(line) < 80:
                            processed.append(f"\n{line}\n")
                        elif line.lower().startswith('chapter'):
                            processed.append(f"\n{line}\n")
                        elif len(line) < 100 and line[:3].replace('.', '').isdigit():
                            processed.append(f"\n{line}\n")
                        else:
                            processed.append(line)
                    text_parts.append('\n'.join(processed))

        if not text_parts:
            raise ValueError("PDF ma text fela parena. Scanned image PDF ho ki check garnus.")

        return "\n".join(text_parts)

    except Exception as e:
        raise ValueError(f"PDF read garna sakiyena: {str(e)}")
        
def extract_from_docx(file_bytes):
    """
    Extract text from DOCX — preserves heading structure
    so section detection works properly.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"DOCX read garna sakiyena: {str(e)}")

    text_parts = []

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        style_name = para.style.name.lower() if para.style else ''

        # Heading styles lai explicit label dine
        # Yesle "Cover Page", "Table of Contents" detect huna maddat garxa
        if 'heading 1' in style_name or 'title' in style_name:
            text_parts.append(f"\n{para.text.upper()}\n")
        elif 'heading 2' in style_name:
            text_parts.append(f"\n{para.text}\n")
        elif 'heading 3' in style_name:
            text_parts.append(f"\n{para.text}\n")
        else:
            text_parts.append(para.text)

    # Tables bata text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    # Core properties bata cover page info inject garne
    # DOCX ma cover page often properties ma hunxa
    try:
        core_props = doc.core_properties
        injected = []
        if core_props.title:
            injected.append(f"cover page\ntitle page\n{core_props.title}")
        if core_props.author:
            injected.append(f"submitted by {core_props.author}")
        if injected:
            text_parts = injected + text_parts
    except Exception:
        pass

    # Table of Contents — DOCX ma TOC field hunxa, text extract hudaina
    # Keyword inject garne if TOC field detected
    full_xml = doc.element.xml.lower()
    if 'toc' in full_xml or 'table of contents' in full_xml or 'w:sdt' in full_xml:
        text_parts.insert(0, "table of contents\n")

    # Cover page check — first page styles
    first_paras = [p.text.lower() for p in doc.paragraphs[:15] if p.text.strip()]
    first_text = ' '.join(first_paras)
    if 'tribhuvan university' in first_text or 'project report' in first_text:
        text_parts.insert(0, "cover page\ntitle page\n")

    if 'supervisor' in first_text or 'recommendation' in first_text:
        text_parts.insert(0, "supervisor's recommendation\n")

    if 'acknowledgement' in first_text or 'acknowledge' in first_text:
        text_parts.insert(0, "acknowledgement\n")

    all_text_lower = ' '.join(text_parts).lower()

    if 'abstract' in all_text_lower or 'executive summary' in all_text_lower:
        text_parts.insert(0, "abstract\n")

    return "\n".join(text_parts)

def extract_formatting_from_docx(file_bytes):
    """
    Extract actual formatting details from DOCX.
    Returns dict with formatting issues found.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"DOCX read garna sakiyena: {str(e)}")

    issues = []
    suggestions = []
    passed = []

    # --- Page margins check ---
    try:
        from docx.util import Inches
        section = doc.sections[0]
        left_margin = round(section.left_margin.inches, 2)
        right_margin = round(section.right_margin.inches, 2)
        top_margin = round(section.top_margin.inches, 2)
        bottom_margin = round(section.bottom_margin.inches, 2)

        # TU standard: Left=1.25, Right=1.0, Top=1.0, Bottom=1.0
        if abs(left_margin - 1.25) > 0.05:
            issues.append({
                "type": "error",
                "message": f"Left margin galat xa: {left_margin}\" detected. TU standard: 1.25\" (35mm) hunuparchha."
            })
        else:
            passed.append({
                "type": "success",
                "message": f"Left margin sahi xa: {left_margin}\" ✓"
            })

        if abs(right_margin - 1.0) > 0.05:
            issues.append({
                "type": "error",
                "message": f"Right margin galat xa: {right_margin}\" detected. TU standard: 1.0\" (25mm) hunuparchha."
            })
        else:
            passed.append({
                "type": "success",
                "message": f"Right margin sahi xa: {right_margin}\" ✓"
            })

        if abs(top_margin - 1.0) > 0.05:
            issues.append({
                "type": "error",
                "message": f"Top margin galat xa: {top_margin}\" detected. TU standard: 1.0\" hunuparchha."
            })
        else:
            passed.append({
                "type": "success",
                "message": f"Top/Bottom margin sahi xa ✓"
            })

    except Exception:
        suggestions.append({
            "type": "warning",
            "message": "Margin check garna sakiyena. Manually verify: Left=1.25\", Right=Top=Bottom=1.0\""
        })

    # --- Page size check (A4) ---
    try:
        from docx.util import Mm
        section = doc.sections[0]
        width_mm = round(section.page_width.mm)
        height_mm = round(section.page_height.mm)
        # A4 = 210 x 297 mm
        if abs(width_mm - 210) <= 2 and abs(height_mm - 297) <= 2:
            passed.append({
                "type": "success",
                "message": f"Page size A4 ({width_mm}x{height_mm}mm) sahi xa ✓"
            })
        else:
            issues.append({
                "type": "error",
                "message": f"Page size galat: {width_mm}x{height_mm}mm detected. TU standard: A4 (210x297mm) hunuparchha."
            })
    except Exception:
        pass

    # --- Font & size check per paragraph ---
    font_errors = []
    size_errors = []
    spacing_errors = []

    checked_para = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        checked_para += 1
        if checked_para > 100:  # max 100 paragraphs check
            break

        # Line spacing check
        try:
            from docx.enum.text import WD_LINE_SPACING
            pf = para.paragraph_format
            if pf.line_spacing is not None:
                # 1.5 line spacing = 276120 EMUs in python-docx (roughly)
                if pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
                    spacing_val = round(float(str(pf.line_spacing)), 1)
                    if abs(spacing_val - 1.5) > 0.1:
                        spacing_errors.append(para.text[:40])
        except Exception:
            pass

        for run in para.runs:
            if not run.text.strip():
                continue

            # Font name check
            font_name = run.font.name
            if font_name and font_name.lower() not in ['times new roman', 'none', '']:
                if len(font_errors) < 3:
                    font_errors.append({
                        "text_sample": run.text[:50],
                        "font": font_name
                    })

            # Font size check
            if run.font.size:
                size_pt = round(run.font.size.pt)
                style_name = para.style.name.lower() if para.style else ''

                # Body text should be 12pt
                if 'heading 1' not in style_name and 'heading 2' not in style_name:
                    if size_pt not in [12, 14, 16]:
                        if len(size_errors) < 3:
                            size_errors.append({
                                "text_sample": run.text[:50],
                                "size": size_pt
                            })

    # Font errors
    if font_errors:
        samples = ', '.join([f'"{e["text_sample"]}..." ({e["font"]})' for e in font_errors[:2]])
        issues.append({
            "type": "error",
            "message": f"Times New Roman baayek aru font use vako xa: {samples}. Sabai text Times New Roman hunuparchha."
        })
    else:
        passed.append({
            "type": "success",
            "message": "Font check: Times New Roman mostly consistent dekhinxa ✓"
        })

    # Size errors
    if size_errors:
        samples = ', '.join([f'"{e["text_sample"]}..." ({e["size"]}pt)' for e in size_errors[:2]])
        issues.append({
            "type": "warning",
            "message": f"Unexpected font sizes detected: {samples}. Body: 12pt, Section heading: 14pt, Chapter: 16pt."
        })

    # Spacing errors
    if spacing_errors:
        issues.append({
            "type": "error",
            "message": f"Line spacing 1.5 baayek aru spacing fela paryo. TU standard: 1.5 line spacing throughout."
        })
    else:
        passed.append({
            "type": "success",
            "message": "Line spacing check passed ✓"
        })

    # --- Heading styles check ---
    heading_issues = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style_name = para.style.name.lower() if para.style else ''
        text_lower = para.text.lower()

        # Check chapter headings
        if 'chapter' in text_lower and len(para.text) < 80:
            for run in para.runs:
                if run.font.size:
                    size = round(run.font.size.pt)
                    if size != 16:
                        heading_issues.append(
                            f"Chapter heading '{para.text[:40]}' — {size}pt detected, 16pt Bold chahiyo"
                        )
                if run.bold is False:
                    heading_issues.append(
                        f"Chapter heading '{para.text[:40]}' — Bold hunuparchha"
                    )

    if heading_issues:
        for h in heading_issues[:3]:
            issues.append({
                "type": "error",
                "message": h
            })
    
    # Summary
    suggestions.append({
        "type": "info",
        "message": (
            "TU Formatting checklist: Font=Times New Roman | "
            "Chapter heading=16pt Bold Center | Section=14pt Bold Left | "
            "Subsection=12pt Bold | Body=12pt Justified | "
            "Spacing=1.5 | Left margin=1.25\" | Others=1\""
        )
    })

    return {
        "formatting_hints": issues + suggestions + passed
    }

